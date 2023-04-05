from docx import Document
import pandas as pd
document = Document('[TIBA] teste.docx')

lista_tags= [str(tag).split(' ')[0] for tag in document.element.body if (str(tag).split(' ')[0] == '<CT_Tbl') or  (str(tag).split(' ')[0] == '<CT_P')]
# print (lista_tags)

tables = []
for table in document.tables:
    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text:
                df[i][j] = cell.text
    tables.append(pd.DataFrame(df))

parrafos = []
for parrafo in document.paragraphs:
    parrafos.append(parrafo.text)


i_tabla = 0
i_parrafo = 0
for elemento in lista_tags:
    if elemento == '<CT_Tbl':
        print(tables[i_tabla])
        i_tabla+=1
    elif str(elemento).split(' ')[0] == '<CT_P':
        print(document.paragraphs[i_parrafo].text)
        i_parrafo+=1
        
    else:
        print(str(elemento))
